import fitz  # PyMuPDF
from docx import Document
from docx.shared import Pt
from docx.section import Section
from docx.text.paragraph import Paragraph
from docx.table import Table
import os
from config.settings import BASE_DIR
import mimetypes
import logging

logger = logging.getLogger(__name__)

class DocumentReaderError(Exception):
    """Exceção customizada para erros de leitura de documento"""
    pass

class DocumentReader:
    @staticmethod
    def validate_file(file_path):
        """Valida o tipo do arquivo usando mimetypes"""
        if not os.path.exists(file_path):
            raise DocumentReaderError("Arquivo não encontrado")
            
        mimetypes.init()
        file_type, _ = mimetypes.guess_type(file_path)
        
        valid_types = {
            'application/pdf': '.pdf',
            'application/vnd.openxmlformats-officedocument.wordprocessingml.document': '.docx'
        }
        
        if not file_type or file_type not in valid_types:
            raise DocumentReaderError(f"Formato de arquivo não suportado: {os.path.splitext(file_path)[1]}")
            
        return file_type

    @staticmethod
    def extract_header_from_pdf(doc, first_page):
        """Extrai o cabeçalho do PDF como imagem fiel ao original"""
        import io
        from PIL import Image
        import uuid

        try:
            # Definir proporção do cabeçalho (ex: 25% do topo)
            header_height_ratio = 0.25
            header_height = int(first_page.rect.height * header_height_ratio)
            header_rect = fitz.Rect(0, 0, first_page.rect.width, header_height)

            # Renderizar a área do cabeçalho como imagem
            header_pix = first_page.get_pixmap(clip=header_rect, dpi=300)
            temp_dir = os.path.join(BASE_DIR, "output", "temp")
            os.makedirs(temp_dir, exist_ok=True)
            header_img_path = os.path.join(temp_dir, f"header_import_{uuid.uuid4().hex}.png")
            header_pix.save(header_img_path)
            return header_img_path  # Retorna o caminho da imagem do header
        except Exception as e:
            logger.error(f"Erro ao extrair cabeçalho do PDF como imagem: {str(e)}")
            return None

    @staticmethod
    def extract_header_from_docx(doc):
        """Extrai cabeçalho do DOCX com máxima fidelidade"""
        header_content = []
        header_images = []
        all_images = []
        image_sizes = {}
        
        try:
            # Extrair cabeçalho do documento
            for section in doc.sections:
                if section.header and section.header.paragraphs:
                    # Processar cada parágrafo do cabeçalho
                    for para in section.header.paragraphs:
                        if not para.text.strip():
                            continue
                            
                        para_data = {
                            'text': para.text,
                            'runs': []
                        }
                        
                        # Processar cada run (segmento com formatação única) do parágrafo
                        for run in para.runs:
                            if not run.text.strip():
                                continue
                                
                            # Capturar todos os detalhes de formatação
                            run_data = {
                                'text': run.text,
                                'style': {
                                    'bold': run.bold,
                                    'italic': run.italic,
                                    'underline': run.underline,
                                },
                                'font': run.font.name if run.font and run.font.name else None,
                                'size': run.font.size if run.font and run.font.size else None,
                                'color': run.font.color.rgb if run.font and run.font.color and run.font.color.rgb else None
                            }
                            
                            para_data['runs'].append(run_data)
                        
                        header_content.append(para_data)
                        
            # Extrair imagens do cabeçalho usando relacionamentos
            temp_dir = os.path.join(BASE_DIR, "output", "temp")
            os.makedirs(temp_dir, exist_ok=True)
            
            # Importar bibliotecas necessárias
            import zipfile
            from io import BytesIO
            from PIL import Image
            
            # Abrir o arquivo docx como um arquivo zip
            with zipfile.ZipFile(BytesIO(doc._element.part.blob)) as zip_file:
                # Extrair imagens do cabeçalho
                for section in doc.sections:
                    if section.header:
                        header_part = section.header.part
                        
                        # Processar relacionamentos de imagem
                        for rel_id, rel in header_part.rels.items():
                            if "image" in rel.reltype:
                                # Obter o caminho da imagem dentro do arquivo zip
                                img_path_in_zip = rel.target_ref
                                if not img_path_in_zip.startswith('/'):
                                    img_path_in_zip = f'word/{img_path_in_zip}'
                                
                                try:
                                    # Extrair a imagem
                                    img_data = zip_file.read(img_path_in_zip.lstrip('/'))
                                    img_filename = os.path.basename(img_path_in_zip)
                                    img_path = os.path.join(temp_dir, f"docx_header_{img_filename}")
                                    
                                    # Salvar a imagem
                                    with open(img_path, 'wb') as img_file:
                                        img_file.write(img_data)
                                    
                                    # Obter dimensões da imagem
                                    with Image.open(img_path) as img:
                                        width, height = img.size
                                    
                                    # Adicionar à lista de imagens do cabeçalho
                                    header_images.append(img_path)
                                    all_images.append(img_path)
                                    
                                    # Estimar posição (sem posição exata disponível no docx)
                                    # Estamos centralizando horizontalmente
                                    image_sizes[img_path] = {
                                        'x': 105 - (width * 0.265) / 2,  # Centralizado (A4 = 210mm / 2 = 105mm)
                                        'y': 20,  # Posição Y estimada no topo
                                        'width': width * 0.265,  # Convertendo pixels para mm (aproximado)
                                        'height': height * 0.265
                                    }
                                
                                except Exception as e:
                                    logger.error(f"Erro ao extrair imagem de documento DOCX: {e}")
        
        except Exception as e:
            logger.error(f"Erro ao extrair cabeçalho do DOCX: {e}")
        
        return header_content, header_images, all_images, image_sizes

    @staticmethod
    def read_pdf(pdf_path):
        """Lê um arquivo PDF e extrai seu conteúdo com máxima fidelidade"""
        import fitz  # PyMuPDF
        
        resultado = {
            'text': '',
            'blocks': [],
            'tables': [],
            'header_images': [],
            'all_images': [],
            'image_sizes': {},
            'colors': {},
            'line_breaks': {},
            'font_info': {},
            'exact_text_positions': [],
            'preserve_original_header': True,
            'original_header_content': [],
            'filename': pdf_path  # Adicionar o caminho do arquivo original
        }
        
        try:
            # Tentar abrir o PDF com opção de reparo
            doc = fitz.open(pdf_path)
            
            # Verificar se o documento foi aberto corretamente
            if not doc or doc.is_closed or doc.page_count == 0:
                logger.error(f"Não foi possível abrir o documento PDF: {pdf_path}")
                raise DocumentReaderError("Não foi possível abrir o documento PDF")
                
            first_page = doc[0]  # Primeira página
            
            # Extrair cabeçalho com máxima fidelidade
            header_img_path = DocumentReader.extract_header_from_pdf(doc, first_page)
            
            resultado['header_images'] = [header_img_path] if header_img_path else []
            
            # Extrair o conteúdo completo do documento
            text_blocks = []
            
            # Processamos página por página
            for page_num in range(len(doc)):
                try:
                    page = doc[page_num]
                    
                    # Extrair texto com formatação
                    try:
                        page_dict = page.get_text("dict")
                    except Exception as text_err:
                        logger.error(f"Erro ao extrair texto da página {page_num}: {str(text_err)}")
                        page_dict = {"blocks": []}
                    
                    for block in page_dict.get("blocks", []):
                        if block.get("type", -1) == 0:  # Bloco de texto
                            block_text = ""
                            font_info = []
                            
                            # Processar linha a linha
                            for line in block.get("lines", []):
                                line_text = ""
                                line_font_info = []
                                
                                # Processar span a span (cada span tem sua própria formatação)
                                for span in line.get("spans", []):
                                    span_text = span.get("text", "").strip()
                                    if span_text:
                                        line_text += span_text
                                        
                                        # Salvar informações de formatação
                                        line_font_info.append({
                                            'font': span.get("font", "Arial"),
                                            'size': span.get("size", 12),
                                            'flags': span.get("flags", 0),
                                            'color': span.get("color", 0),
                                            'style': {
                                                'bold': bool(span.get("flags", 0) & 1),
                                                'italic': bool(span.get("flags", 0) & 2),
                                                'underline': bool(span.get("flags", 0) & 4)
                                            }
                                        })
                                
                                # Adicionar quebra de linha após cada linha
                                if line_text:
                                    block_text += line_text + "\n"
                                    font_info.extend(line_font_info)
                            
                            # Remover quebra de linha extra no final
                            block_text = block_text.rstrip()
                            
                            if block_text:
                                # Verificar se este bloco é uma tabela com algoritmo melhorado
                                is_table = False
                                
                                # Heurística aprimorada para detecção de tabelas:
                                # 1. Detecção por padrões de formatação (tabulações, separadores)
                                # 2. Análise de estrutura de linhas e colunas
                                # 3. Detecção de bordas
                                
                                # Verificar padrões tabulares clássicos
                                tab_patterns = ["\t", "|", "  ", "—", "–", "+", "-+", "+---", "----"]
                                grid_patterns = ["+-+", "+--+", "|--|", "┌", "┐", "└", "┘", "├", "┤", "┬", "┴", "┼", "│", "─"]
                                
                                # Verificar estrutura e padrões
                                lines = block_text.split("\n")
                                
                                # Avaliar alinhamento de caracteres em múltiplas linhas (indica tabela)
                                if len(lines) > 1:
                                    char_positions = []
                                    alignment_score = 0
                                    
                                    # Coletar posições dos caracteres específicos que podem indicar colunas
                                    for line in lines:
                                        positions = [i for i, char in enumerate(line) if char in ":|.-+"]
                                        if positions:
                                            char_positions.append(positions)
                                    
                                    # Verificar alinhamentos verticais (colunas)
                                    if len(char_positions) > 1:
                                        all_positions = set()
                                        for positions in char_positions:
                                            all_positions.update(positions)
                                        
                                        # Para cada posição, verificar quantas linhas têm um caractere nela
                                        for pos in all_positions:
                                            aligned_count = sum(1 for positions in char_positions if pos in positions)
                                            if aligned_count > len(char_positions) // 2:
                                                alignment_score += 1
                                    
                                    # Se temos vários alinhamentos verticais, provavelmente é uma tabela
                                    if alignment_score >= 2:
                                        is_table = True
                                
                                # Verificar presença de padrões de tabela
                                if not is_table:
                                    # Verificar padrões comuns de grid
                                    for pattern in grid_patterns:
                                        if pattern in block_text:
                                            is_table = True
                                            break
                                    
                                    # Verificar padrões básicos indicando estrutura tabular
                                    if not is_table and (
                                        "\t" in block_text or  # Tabulações são um forte indicador
                                        (block_text.count("\n") > 1 and  # Múltiplas linhas com padrões
                                         any(pattern in block_text for pattern in tab_patterns))
                                    ):
                                        # Verificar consistência entre linhas (mesma quantidade de separadores)
                                        if len(lines) > 1:
                                            # Contar separadores em cada linha
                                            separators_count = []
                                            for line in lines:
                                                if line.strip():  # Ignorar linhas vazias
                                                    count = (line.count("|") + 
                                                             line.count("\t") + 
                                                             line.count("  "))
                                                    separators_count.append(count)
                                            
                                            # Se a maioria das linhas tem número semelhante de separadores, é uma tabela
                                            if separators_count and len(set(separators_count)) <= 2:  # Permitir variação de 1
                                                is_table = True
                                
                                # Verificar também análise de fonte monospace e posicionamento pixel-perfect
                                if not is_table and block.get("font_info"):
                                    # Verificar se usa fonte monospace (comum em tabelas)
                                    monospace_fonts = ["Courier", "Consolas", "Monaco", "Menlo", "MonoSpace"]
                                    for font_info in block.get("font_info", []):
                                        if any(mono in font_info.get("font", "") for mono in monospace_fonts):
                                            is_table = True
                                            break
                                
                                # Verificar caracteres de bordas internacionais
                                border_chars = "┌┐└┘├┤┬┴┼│─"
                                if not is_table and any(char in block_text for char in border_chars):
                                    is_table = True
                                
                                # Adicionar à lista de tabelas se identificado como tabela
                                if is_table:
                                    # Melhorar a estrutura de dados da tabela para preservar a formatação
                                    tabela_estruturada = {
                                        'text': block_text,
                                        'bbox': block.get("bbox", [0, 0, 0, 0]),
                                        'font_info': font_info,
                                        'linhas': lines,
                                        'estrutura_detectada': True
                                    }
                                    resultado['tables'].append(tabela_estruturada)
                                
                                # Adicionar bloco ao resultado
                                text_block = {
                                    'text': block_text,
                                    'bbox': block.get("bbox", [0, 0, 0, 0]),
                                    'is_table': is_table,
                                    'font_info': font_info
                                }
                                
                                text_blocks.append(text_block)
                        
                        elif block.get("type", -1) == 1:  # Bloco de imagem
                            # As imagens já foram tratadas anteriormente
                            pass
                            
                except Exception as page_err:
                    logger.error(f"Erro ao processar página {page_num}: {str(page_err)}")
                    continue
            
            # Montar texto completo e armazenar blocos
            full_text = "\n\n".join(block.get('text', '') for block in text_blocks)
            resultado['text'] = full_text
            resultado['blocks'] = text_blocks
            
            doc.close()
        
        except Exception as e:
            error_msg = f"Erro ao ler PDF: {str(e)}"
            logger.error(error_msg)
            raise DocumentReaderError(error_msg)
        
        return resultado

    @staticmethod
    def read_docx(docx_path):
        """Lê um arquivo DOCX e extrai seu conteúdo com máxima fidelidade"""
        from docx import Document
        
        resultado = {
            'text': '',
            'blocks': [],
            'tables': [],
            'header_images': [],
            'all_images': [],
            'image_sizes': {},
            'colors': {},
            'line_breaks': {},
            'font_info': {},
            'preserve_original_header': True,
            'original_header_content': [],
            'filename': docx_path  # Adicionar o caminho do arquivo original
        }
        
        try:
            doc = Document(docx_path)
            
            # Extrair cabeçalho com máxima fidelidade
            header_content, header_images, all_images, image_sizes = DocumentReader.extract_header_from_docx(doc)
            
            resultado['original_header_content'] = header_content
            resultado['header_images'] = header_images
            resultado['all_images'] = all_images
            resultado['image_sizes'] = image_sizes
            
            # Processar conteúdo do documento
            text_blocks = []
            
            # Processar parágrafos
            for para in doc.paragraphs:
                if not para.text.strip():
                    continue
                
                # Informações do parágrafo
                para_data = {
                    'text': para.text,
                    'is_table': False,
                    'font_info': []
                }
                
                # Extrair informações de formatação de cada run
                for run in para.runs:
                    if not run.text.strip():
                        continue
                    
                    # Capturar informações completas de formatação
                    font_info = {
                        'font': run.font.name if run.font and run.font.name else "Arial",
                        'size': run.font.size if run.font and run.font.size else 12,
                        'style': {
                            'bold': run.bold if run.bold is not None else False,
                            'italic': run.italic if run.italic is not None else False,
                            'underline': run.underline if run.underline is not None else False
                        },
                        'color': run.font.color.rgb if run.font and run.font.color and run.font.color.rgb else None
                    }
                    
                    para_data['font_info'].append(font_info)
                
                text_blocks.append(para_data)
            
            # Processar tabelas
            for table in doc.tables:
                table_data = {
                    'text': '',
                    'is_table': True,
                    'font_info': []
                }
                
                # Converter tabela em texto para processamento posterior
                for i, row in enumerate(table.rows):
                    for j, cell in enumerate(row.cells):
                        cell_text = cell.text.strip()
                        if cell_text:
                            # Coletar informações de formatação da célula
                            for paragraph in cell.paragraphs:
                                for run in paragraph.runs:
                                    font_info = {
                                        'font': run.font.name if run.font and run.font.name else "Arial",
                                        'size': run.font.size if run.font and run.font.size else 12,
                                        'style': {
                                            'bold': run.bold if run.bold is not None else False,
                                            'italic': run.italic if run.italic is not None else False,
                                            'underline': run.underline if run.underline is not None else False
                                        }
                                    }
                                    table_data['font_info'].append(font_info)
                            
                            # Adicionar o texto formatado da célula
                            table_data['text'] += f"{cell_text} | "
                    
                    # Quebra de linha no fim de cada linha da tabela
                    table_data['text'] += "\n"
                
                # Remover caracteres extras e adicionar a tabela aos blocos
                table_data['text'] = table_data['text'].replace(" | \n", "\n").rstrip(" |\n")
                if table_data['text'].strip():
                    text_blocks.append(table_data)
                    resultado['tables'].append(table_data)
            
            # Montar texto completo e adicionar blocos
            full_text = "\n\n".join(block['text'] for block in text_blocks)
            resultado['text'] = full_text
            resultado['blocks'] = text_blocks
        
        except Exception as e:
            error_msg = f"Erro ao ler DOCX: {str(e)}"
            logger.error(error_msg)
            raise DocumentReaderError(error_msg)
        
        return resultado

    @staticmethod
    def read(file_path):
        """Método principal para ler documentos, detectando automaticamente o tipo"""
        ext = os.path.splitext(file_path)[1].lower()
        
        if ext == '.pdf':
            return DocumentReader.read_pdf(file_path)
        elif ext in ['.docx', '.doc']:
            return DocumentReader.read_docx(file_path)
        else:
            raise DocumentReaderError("Formato de arquivo não suportado")
