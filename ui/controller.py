from core import reader, randomizer, generator, ai_helper
from core.pdf_viewer import PDFHeaderViewer
from PyQt6 import uic, QtWidgets, QtGui
from PyQt6.QtCore import Qt, QTimer
from PyQt6.QtWidgets import QDialog, QVBoxLayout, QHBoxLayout, QLabel, QPushButton, QApplication
from PyQt6.QtGui import QIcon, QColor, QPixmap
from config.settings import BASE_DIR, ASSETS_DIR, GEMINI_API_KEY
import os
import json

class ErrorDialog(QDialog):
    def __init__(self, title, message, parent=None):
        super().__init__(parent)
        self.setWindowTitle(title)
        self.setMinimumWidth(400)
        
        layout = QVBoxLayout()
        
        # Área da mensagem
        message_layout = QHBoxLayout()
        icon_label = QLabel()
        icon_label.setPixmap(self.style().standardIcon(self.style().StandardPixmap.SP_MessageBoxCritical).pixmap(32, 32))
        message_layout.addWidget(icon_label)
        message_label = QLabel(message)
        message_label.setWordWrap(True)
        message_layout.addWidget(message_label, 1)
        layout.addLayout(message_layout)
        
        # Botões
        button_layout = QHBoxLayout()
        self.copy_button = QPushButton("Copiar Erro")  # Store as instance variable
        self.copy_button.clicked.connect(lambda: self.copy_to_clipboard(message))
        ok_button = QPushButton("OK")
        ok_button.clicked.connect(self.accept)
        
        button_layout.addWidget(self.copy_button)
        button_layout.addWidget(ok_button)
        layout.addLayout(button_layout)
        
        self.setLayout(layout)
    
    def copy_to_clipboard(self, text):
        QApplication.clipboard().setText(text)
        # Feedback visual temporário usando referência direta ao botão
        self.copy_button.setText("Copiado!")
        QTimer.singleShot(1000, lambda: self.copy_button.setText("Copiar Erro"))

class ProvaController:
    def __init__(self, ui):
        self.ui = ui
        self.formato_original = None

    def carregar_prova(self, caminho):
        if caminho.endswith(".pdf"):
            resultado = reader.DocumentReader.read_pdf(caminho)
        elif caminho.endswith(".docx"):
            resultado = reader.DocumentReader.read_docx(caminho)
        else:
            raise Exception("Formato não suportado")
            
        self.formato_original = resultado
        return self.processar_questoes(resultado['text'])

    def processar_questoes(self, texto):
        """Processa as questões mantendo a formatação original"""
        questoes = []
        blocos = self.formato_original.get('blocks', [])
        
        questao_atual = None
        alternativas_atuais = []
        tem_numeracao_propria = False
        tem_palavra_questao = False
        
        for bloco in blocos:
            texto_bloco = bloco.get('text', '').strip()
            if not texto_bloco:
                continue
            
            # Verifica se o texto já contém numeração e a palavra "Questão"
            if questao_atual is None and any(texto_bloco.startswith(f"Questão {i}") for i in range(1, 100)):
                tem_palavra_questao = True
            if questao_atual is None and any(texto_bloco.startswith(f"{i}.") or texto_bloco.startswith(f"{i})") for i in range(1, 100)):
                tem_numeracao_propria = True
                
            # Verificar se é início de questão
            if (any(texto_bloco.startswith(f"{i}.") or texto_bloco.startswith(f"{i})") for i in range(1, 100)) or
                any(texto_bloco.startswith(f"Questão {i}") for i in range(1, 100))):
                if questao_atual is not None:
                    questoes.append((questao_atual, alternativas_atuais))
                    alternativas_atuais = []
                questao_atual = texto_bloco
                
            # Verificar se é alternativa
            elif any(texto_bloco.startswith(f"({letra})") or texto_bloco.startswith(f"{letra})") 
                    for letra in "ABCDEabcde"):
                # Remove possíveis duplicações da letra da alternativa
                for letra in "ABCDEabcde":
                    if texto_bloco.startswith(f"({letra})"):
                        texto_bloco = f"({letra})" + texto_bloco.replace(f"({letra})", "", 1).strip()
                        break
                    elif texto_bloco.startswith(f"{letra})"):
                        texto_bloco = f"{letra})" + texto_bloco.replace(f"{letra})", "", 1).strip()
                        break
                
                if questao_atual is not None:
                    alternativas_atuais.append(texto_bloco)
            
            # Se é tabela, incluir na questão atual com formatação aprimorada
            elif bloco.get('is_table', False):
                if questao_atual is not None:
                    # Verificar se temos informações detalhadas da tabela no formato original
                    tabela_detalhada = None
                    for tabela in self.formato_original.get('tables', []):
                        if tabela.get('text', '') == texto_bloco:
                            tabela_detalhada = tabela
                            break
                    
                    # Se temos informações estruturadas da tabela, usar para renderização mais precisa
                    if tabela_detalhada and tabela_detalhada.get('estrutura_detectada', False):
                        linhas = tabela_detalhada.get('linhas', texto_bloco.split('\n'))
                        
                        # Construir HTML da tabela
                        tabela_html = "<table style='border-collapse: collapse; width: 100%; margin: 10px 0;'>"
                        
                        # Analisar a estrutura para detectar cabeçalho
                        tem_cabecalho = False
                        if len(linhas) > 1:
                            primeira_linha = linhas[0]
                            segunda_linha = linhas[1]
                            if segunda_linha and all(c in '-+=' for c in segunda_linha if c.strip()):
                                tem_cabecalho = True
                        
                        # Processar cada linha da tabela
                        linha_atual = 0
                        while linha_atual < len(linhas):
                            linha = linhas[linha_atual]
                            
                            # Pular linhas de separação
                            if linha and all(c in '-+=' for c in linha if c.strip()):
                                linha_atual += 1
                                continue
                            
                            # Começar nova linha na tabela HTML
                            tabela_html += "<tr>"
                            
                            # Determinar se é linha de cabeçalho
                            celula_tag = "th" if (tem_cabecalho and linha_atual == 0) else "td"
                            estilo_celula = "style='border: 1px solid #ddd; padding: 8px; text-align: left;"
                            if celula_tag == "th":
                                estilo_celula += " background-color: #f2f2f2; font-weight: bold;'"
                            else:
                                estilo_celula += "'"
                            
                            # Detectar colunas usando delimitadores comuns
                            colunas = []
                            if "|" in linha:
                                # Separar por pipe
                                colunas = [col.strip() for col in linha.split("|")]
                                # Remover células vazias nas extremidades
                                if colunas and not colunas[0].strip():
                                    colunas.pop(0)
                                if colunas and not colunas[-1].strip():
                                    colunas.pop()
                            elif "\t" in linha:
                                # Separar por tabulação
                                colunas = [col.strip() for col in linha.split("\t")]
                            elif "  " in linha:
                                # Tentar separar por múltiplos espaços (mais difícil)
                                import re
                                colunas = [col.strip() for col in re.split(r'\s{2,}', linha)]
                            else:
                                # Se não identificou colunas, tratar como célula única
                                colunas = [linha]
                            
                            # Adicionar cada coluna como célula
                            for coluna in colunas:
                                if coluna.strip():
                                    tabela_html += f"<{celula_tag} {estilo_celula}>{coluna}</{celula_tag}>"
                            
                            tabela_html += "</tr>"
                            linha_atual += 1
                        
                        tabela_html += "</table>"
                        questao_atual += f"\n{tabela_html}"
                    else:
                        # Fallback - converter para tabela simples com separação por pipe
                        linhas = texto_bloco.split('\n')
                        tabela_html = "<table style='border-collapse: collapse; width: 100%; margin: 10px 0;'>"
                        
                        for i, linha in enumerate(linhas):
                            if not linha.strip():
                                continue
                            
                            # Determinar se é linha de cabeçalho
                            e_cabecalho = (i == 0)
                            celula_tag = "th" if e_cabecalho else "td"
                            estilo_celula = "style='border: 1px solid #ddd; padding: 8px; text-align: left;"
                            if celula_tag == "th":
                                estilo_celula += " background-color: #f2f2f2; font-weight: bold;'"
                            else:
                                estilo_celula += "'"
                            
                            # Detectar colunas usando delimitadores comuns
                            colunas = []
                            if "|" in linha:
                                colunas = [col.strip() for col in linha.split("|")]
                            elif "\t" in linha:
                                colunas = [col.strip() for col in linha.split("\t")]
                            else:
                                # Se não tem separadores claros, tentar outros padrões
                                colunas = [linha]
                            
                            # Adicionar linha à tabela
                            tabela_html += "<tr>"
                            for coluna in colunas:
                                if coluna.strip():
                                    tabela_html += f"<{celula_tag} {estilo_celula}>{coluna}</{celula_tag}>"
                            tabela_html += "</tr>"
                        
                        tabela_html += "</table>"
                        questao_atual += f"\n{tabela_html}"
        
        # Adicionar última questão se houver
        if questao_atual is not None and alternativas_atuais:
            questoes.append((questao_atual, alternativas_atuais))
        
        # Guardar informações sobre o formato original
        self.formato_original['tem_numeracao_propria'] = tem_numeracao_propria
        self.formato_original['tem_palavra_questao'] = tem_palavra_questao
        
        return questoes

    def aplicar_embaralhamento(self, questoes, modo):
        if modo == "questoes":
            return randomizer.embaralhar_questoes(questoes)
        elif modo == "alternativas":
            return [randomizer.embaralhar_alternativas(q) for q in questoes]
        elif modo == "tudo":
            return randomizer.embaralhar_tudo(questoes)
        else:
            return questoes

    def aplicar_ia(self, questoes, api_key):
        novas = []
        for enunciado, alternativas in questoes:
            novo = ai_helper.substituir_por_sinonimos(enunciado, api_key)
            novas.append((novo, alternativas))
        return novas

    def gerar_info_prova(self):
        """Gera informações para o QR Code"""
        return {
            'titulo': 'Prova gerada com ProvaGuard',
            'modo_embaralhamento': self.ui.shuffleCombo.currentText(),
            'ia_utilizada': self.ui.aiCheckBox.isChecked()
        }

class MainController(QtWidgets.QMainWindow):
    def __init__(self):
        super().__init__()
        # Load UI
        ui_file = os.path.join(BASE_DIR, "ui", "main_window.ui")
        uic.loadUi(ui_file, self)
        
        # Load theme
        theme_file = os.path.join(BASE_DIR, "ui", "styles", "theme.qss")
        with open(theme_file, "r") as f:
            self.setStyleSheet(f.read())
        
        # Setup UI elements
        self.setup_ui()
        self.connect_signals()
        
        self.prova_controller = ProvaController(self)
        self.questoes_atuais = None
        
        self.last_qrcode_data = None
        self.last_gabarito_path = None
        
        # Initial state
        self.apiKeyInput.hide()
        self.progressBar.setValue(0)
        self.uploadBtn.setProperty("loaded", False)
        self.style().unpolish(self.uploadBtn)
        self.style().polish(self.uploadBtn)
        
    def setup_ui(self):
        self.setWindowTitle("ProvaGuard")
        icons_dir = os.path.join(BASE_DIR, "ui", "assets", "icons")
        
        try:
            # Configurar ícones usando caminhos absolutos e verificando existência
            icons = {
                'upload': "upload.svg",
                'preview': "preview.svg",
                'generate': "generate.svg",
                'print': "print.svg",
                'dropdown': "dropdown.svg"
            }
            
            for name, file in icons.items():
                icon_path = os.path.join(icons_dir, file)
                if not os.path.exists(icon_path):
                    print(f"Aviso: Ícone não encontrado: {icon_path}")
                    continue
                    
                if name == 'upload':
                    self.uploadBtn.setIcon(QIcon(icon_path))
                elif name == 'preview':
                    self.previewBtn.setIcon(QIcon(icon_path))
                elif name == 'generate':
                    self.generateBtn.setIcon(QIcon(icon_path))
                elif name == 'print':
                    self.printBtn.setIcon(QIcon(icon_path))
                elif name == 'dropdown':
                    self.shuffleCombo.setStyleSheet(f"""
                        QComboBox::down-arrow {{
                            image: url({icon_path.replace(os.sep, '/')});
                            width: 12px;
                            height: 12px;
                        }}
                    """)
            
        except Exception as e:
            print(f"Erro ao carregar ícones: {e}")
        
        # Configurar ComboBox
        self.shuffleCombo.setMaxVisibleItems(10)
        self.shuffleCombo.view().setVerticalScrollBarPolicy(Qt.ScrollBarPolicy.ScrollBarAsNeeded)
        self.shuffleCombo.view().setHorizontalScrollBarPolicy(Qt.ScrollBarPolicy.ScrollBarAlwaysOff)
        
        # Desabilitar botões até carregar prova
        self.previewBtn.setEnabled(False)
        self.generateBtn.setEnabled(False)
        self.printBtn.setEnabled(False)
        self.shuffleCombo.setEnabled(False)
        self.aiCheckBox.setEnabled(False)

        # Adicionar botão de visualizar gabarito
        self.gabarito_btn = QPushButton("Ver Gabarito")
        self.gabarito_btn.setEnabled(False)
        self.buttonLayout.addWidget(self.gabarito_btn)
        
        # Configurar preview do QR Code
        self.qr_preview = QLabel()
        self.qr_preview.setMinimumSize(100, 100)
        self.qr_preview.setMaximumSize(100, 100)
        self.qr_preview.setAlignment(Qt.AlignmentFlag.AlignCenter)
        self.qr_preview.setStyleSheet("background-color: white; border-radius: 5px; padding: 5px;")
        self.mainLayout.insertWidget(self.mainLayout.indexOf(self.progressBar), self.qr_preview)
        
        # Variáveis para cabeçalho e rodapé
        self.header_image_path = None
        self.image_preview = QLabel()
        self.image_preview.setMaximumSize(80, 80)
        self.image_preview.setAlignment(Qt.AlignmentFlag.AlignCenter)
        self.image_preview.setStyleSheet("background-color: white; border-radius: 5px;")
        self.headerFooterLayout.addWidget(self.image_preview)

    def connect_signals(self):
        self.uploadBtn.clicked.connect(self.on_upload)
        self.previewBtn.clicked.connect(self.on_preview)
        self.generateBtn.clicked.connect(self.on_generate)
        self.printBtn.clicked.connect(self.on_print)
        self.aiCheckBox.stateChanged.connect(self.on_ai_toggle)
        self.gabarito_btn.clicked.connect(self.on_view_gabarito)
        # Conectar sinais para cabeçalho e rodapé
        self.headerImageBtn.clicked.connect(self.on_select_header_image)
        self.importLogoBtn.clicked.connect(self.on_import_logo_from_doc)
        self.importHeaderBtn.clicked.connect(self.on_import_header_complete)
    
    def show_error(self, title, message):
        dialog = ErrorDialog(title, str(message), self)
        dialog.exec()
    
    def show_success_message(self, message):
        success_label = QLabel(message)
        success_label.setStyleSheet("""
            QLabel {
                background-color: #4CAF50;
                color: white;
                padding: 10px;
                border-radius: 5px;
                font-weight: bold;
            }
        """)
        success_label.setAlignment(Qt.AlignmentFlag.AlignCenter)
        
        dialog = QDialog(self)
        dialog.setWindowTitle("Sucesso")
        dialog.setWindowFlags(Qt.WindowType.FramelessWindowHint | Qt.WindowType.Popup)
        
        layout = QVBoxLayout()
        layout.addWidget(success_label)
        dialog.setLayout(layout)
        
        # Posicionar no centro da janela principal
        dialog.move(
            self.x() + (self.width() - dialog.width()) // 2,
            self.y() + (self.height() - dialog.height()) // 2
        )
        
        dialog.show()
        QTimer.singleShot(2000, dialog.close)  # Fecha após 2 segundos
    
    def on_upload(self):
        file_path, _ = QtWidgets.QFileDialog.getOpenFileName(
            self,
            "Selecionar arquivo",
            "",
            "Documents (*.pdf *.docx)"
        )
        if file_path:
            try:
                self.questoes_atuais = self.prova_controller.carregar_prova(file_path)
                
                # Atualizar estado do botão
                self.uploadBtn.setProperty("loaded", True)
                self.uploadBtn.setText("Arquivo Carregado ✓")
                self.uploadBtn.setStyleSheet("""
                    QPushButton[loaded="true"] {
                        background-color: #4CAF50;
                    }
                """)
                
                # Habilitar outros controles
                self.previewBtn.setEnabled(True)
                self.generateBtn.setEnabled(True)
                self.printBtn.setEnabled(True)
                self.shuffleCombo.setEnabled(True)
                self.aiCheckBox.setEnabled(True)
                
                # Mostrar mensagem de sucesso
                self.show_success_message("Prova carregada com sucesso!")
                self.statusbar.showMessage(f"Arquivo carregado: {os.path.basename(file_path)}", 3000)
                
            except Exception as e:
                self.show_error("Erro", str(e))
                self.uploadBtn.setProperty("loaded", False)
                self.style().unpolish(self.uploadBtn)
                self.style().polish(self.uploadBtn)
    
    def on_preview(self):
        if not self.questoes_atuais:
            QtWidgets.QMessageBox.warning(self, "Aviso", "Carregue um arquivo primeiro!")
            return
        
        # Criar um diálogo para mostrar o preview do PDF original
        preview_dialog = QtWidgets.QDialog(self)
        preview_dialog.setWindowTitle("Preview da Prova Original")
        preview_dialog.setMinimumSize(800, 600)  # Tamanho reduzido para caber na maioria das telas
        
        layout = QtWidgets.QVBoxLayout()
        
        # Obter o caminho do PDF original
        formato_original = self.prova_controller.formato_original or {}
        original_path = formato_original.get('filename', '')
        
        if original_path and os.path.exists(original_path):
            # Usar o PDFHeaderViewer para mostrar o PDF original sem modificações
            pdf_viewer = PDFHeaderViewer()
            pdf_viewer.load_pdf(original_path)
            layout.addWidget(pdf_viewer)
            
            # Adicionar label explicativo
            help_label = QtWidgets.QLabel("Use os controles para navegar entre páginas e ajustar o zoom. Use a barra de rolagem para ver todo o conteúdo.")
            help_label.setStyleSheet("color: #666; font-style: italic;")
            help_label.setAlignment(Qt.AlignmentFlag.AlignCenter)
            layout.addWidget(help_label)
            
            # Botões
            buttons = QtWidgets.QDialogButtonBox(
                QtWidgets.QDialogButtonBox.StandardButton.Close
            )
            buttons.rejected.connect(preview_dialog.reject)
            layout.addWidget(buttons)
            
            preview_dialog.setLayout(layout)
            
            # Ajustar o diálogo para 80% do tamanho da tela
            screen_size = QtWidgets.QApplication.primaryScreen().size()
            preview_dialog.resize(int(screen_size.width() * 0.8), int(screen_size.height() * 0.8))
            
            preview_dialog.exec()
        else:
            # Se não encontrar o arquivo original, mostrar mensagem de erro
            self.show_error("Erro", "Não foi possível encontrar o arquivo original para preview.")
    
    def on_generate(self):
        if not self.questoes_atuais:
            self.show_error("Aviso", "Carregue um arquivo primeiro!")
            return
            
        questoes = self.questoes_atuais
        modo = self.shuffleCombo.currentText().lower().split()[1]
        
        # Aplicar embaralhamento
        questoes = self.prova_controller.aplicar_embaralhamento(questoes, modo)
        
        # Aplicar IA se selecionado
        if self.aiCheckBox.isChecked():
            questoes = self.prova_controller.aplicar_ia(questoes, GEMINI_API_KEY)
        
        # Gerar QR Code e PDF
        try:
            os.makedirs("output/provas_geradas", exist_ok=True)
            qr_path = "output/provas_geradas/qrcode_temp.png"
            
            # Gerar QR Code com informações da prova
            info_prova = self.prova_controller.gerar_info_prova()
            self.last_qrcode_data = generator.gerar_qrcode(info_prova, qr_path)
            
            # Atualizar preview do QR Code
            pixmap = QPixmap(qr_path)
            self.qr_preview.setPixmap(pixmap.scaled(100, 100, Qt.AspectRatioMode.KeepAspectRatio))
            
            # Preparar informações de cabeçalho e rodapé
            header_footer_info = {
                'header_text': self.headerTextInput.text(),
                'footer_text': self.footerTextInput.text(),
                'header_image': self.header_image_path,
                # Informações adicionais para o cabeçalho escolar
                'teacher_info': self.teacherInput.text(),
                'subject': self.subjectInput.text(),
                'block_info': self.blockInput.text(),
                'evaluation_type': self.evaluationInput.text()
            }
            
            formato_original = self.prova_controller.formato_original or {}
            # Atualizar o formato original com as informações de cabeçalho e rodapé
            if header_footer_info['header_text']:
                formato_original['header_text'] = header_footer_info['header_text']
            if header_footer_info['footer_text']:
                formato_original['footer_text'] = header_footer_info['footer_text']
            if header_footer_info['header_image'] and os.path.exists(header_footer_info['header_image']):
                # Se o formato original já tinha imagens, adicionar esta nova
                if 'header_images' not in formato_original:
                    formato_original['header_images'] = []
                formato_original['header_images'].append(header_footer_info['header_image'])
                
            # Adicionar informações extras para o cabeçalho escolar
            formato_original['teacher_info'] = header_footer_info['teacher_info']
            formato_original['subject'] = header_footer_info['subject']
            formato_original['block_info'] = header_footer_info['block_info']
            formato_original['evaluation_type'] = header_footer_info['evaluation_type']
            
            # Gerar PDF e gabarito
            pdf_path, gabarito_path = generator.gerar_pdf_prova(
                "prova_gerada", 
                questoes, 
                qr_path,
                formato_original
            )
            
            self.last_gabarito_path = gabarito_path
            self.gabarito_btn.setEnabled(True)
            
            self.statusbar.showMessage(f"Prova gerada com sucesso! Salva em: {pdf_path}", 5000)
        except Exception as e:
            self.show_error("Erro", f"Erro ao gerar prova: {str(e)}")
    
    def on_print(self):
        pdf_path = "output/provas_geradas/prova_gerada.pdf"
        if not os.path.exists(pdf_path):
            self.show_error("Aviso", "Gere uma prova primeiro!")
            return
        try:
            from core import printer
            printer.imprimir(pdf_path)
            self.statusbar.showMessage("Prova enviada para impressão!", 3000)
        except Exception as e:
            self.show_error("Erro", f"Erro ao imprimir: {str(e)}")
    
    def on_ai_toggle(self, state):
        self.apiKeyInput.setEnabled(state == Qt.CheckState.Checked.value)
    
    def on_view_gabarito(self):
        """Abre o gabarito gerado"""
        if not self.last_gabarito_path or not os.path.exists(self.last_gabarito_path):
            self.show_error("Aviso", "Gabarito não encontrado! Gere a prova novamente.")
            return
        
        try:
            if os.name == 'nt':  # Windows
                os.startfile(self.last_gabarito_path)
            else:  # Linux/Mac
                import subprocess
                subprocess.run(['xdg-open', self.last_gabarito_path])
        except Exception as e:
            self.show_error("Erro", f"Erro ao abrir gabarito: {str(e)}")

    def on_select_header_image(self):
        """Selecionar imagem para o cabeçalho"""
        file_path, _ = QtWidgets.QFileDialog.getOpenFileName(
            self,
            "Selecionar Logo",
            "",
            "Imagens (*.png *.jpg *.jpeg)"
        )
        
        if file_path:
            self.header_image_path = file_path
            pixmap = QPixmap(file_path)
            # Redimensionar para visualização
            scaled_pixmap = pixmap.scaled(80, 80, Qt.AspectRatioMode.KeepAspectRatio)
            self.image_preview.setPixmap(scaled_pixmap)
            self.selectedImageLabel.setText(os.path.basename(file_path))

    def on_import_logo_from_doc(self):
        """Importar logo diretamente de um arquivo PDF ou DOCX"""
        file_path, _ = QtWidgets.QFileDialog.getOpenFileName(
            self,
            "Selecionar Arquivo com Logo",
            "",
            "Documentos (*.pdf *.docx)"
        )
        
        if not file_path:
            return
            
        try:
            self.statusbar.showMessage(f"Extraindo logo de {os.path.basename(file_path)}...", 2000)
            self.progressBar.setValue(30)
            
            # Diretório temporário para salvar imagens extraídas
            temp_dir = os.path.join(BASE_DIR, "output", "temp")
            os.makedirs(temp_dir, exist_ok=True)
            
            # Extrair logo com base no tipo de arquivo
            images = []
            if file_path.endswith(".pdf"):
                # Extrair imagens do PDF
                import fitz  # PyMuPDF
                try:
                    doc = fitz.open(file_path)
                    page = doc[0]  # Extrair da primeira página
                    
                    # Obter imagens da página
                    for img_index, img in enumerate(page.get_images(full=True)):
                        xref = img[0]
                        if xref:
                            pix = fitz.Pixmap(doc, xref)
                            if pix.width > 0 and pix.height > 0:
                                img_path = os.path.join(temp_dir, f"extracted_logo_{img_index}.png")
                                if pix.alpha:
                                    pix = fitz.Pixmap(fitz.csRGB, pix)
                                pix.save(img_path)
                                images.append(img_path)
                    
                    doc.close()
                except Exception as e:
                    self.show_error("Erro", f"Erro ao extrair imagem do PDF: {str(e)}")
                    self.progressBar.setValue(0)
                    return
                    
            elif file_path.endswith(".docx"):
                # Extrair imagens do DOCX
                import zipfile
                from docx import Document
                from PIL import Image
                import io
                
                try:
                    doc = Document(file_path)
                    
                    # Método 1: Extrair através do relacionamento de imagens
                    zip_file = zipfile.ZipFile(file_path)
                    for rel in doc.part.rels.values():
                        if "image" in rel.reltype:
                            image_name = rel.target_ref.split('/')[-1]
                            image_data = zip_file.read(f"word/{rel.target_ref}")
                            img_path = os.path.join(temp_dir, f"extracted_{image_name}")
                            
                            with open(img_path, "wb") as f:
                                f.write(image_data)
                            
                            images.append(img_path)
                    
                    # Se não encontrou imagens, tentar cabeçalhos
                    if not images:
                        for section in doc.sections:
                            if section.header:
                                header_rels = section.header.part.rels
                                for rel in header_rels.values():
                                    if "image" in rel.reltype:
                                        image_name = rel.target_ref.split('/')[-1]
                                        image_data = zip_file.read(f"word/{rel.target_ref}")
                                        img_path = os.path.join(temp_dir, f"header_{image_name}")
                                        
                                        with open(img_path, "wb") as f:
                                            f.write(image_data)
                                        
                                        images.append(img_path)
                                        
                except Exception as e:
                    self.show_error("Erro", f"Erro ao extrair imagem do DOCX: {str(e)}")
                    self.progressBar.setValue(0)
                    return
            
            self.progressBar.setValue(70)
            
            # Selecionar a primeira imagem encontrada ou mostrar diálogo de seleção
            if not images:
                self.show_error("Aviso", "Nenhuma imagem encontrada no documento!")
                self.progressBar.setValue(0)
                return
            
            if len(images) == 1:
                selected_image = images[0]
            else:
                # Criar diálogo para escolher entre as imagens encontradas
                dialog = QDialog(self)
                dialog.setWindowTitle("Selecione o Logo")
                dialog.setMinimumWidth(600)
                
                layout = QVBoxLayout()
                layout.addWidget(QLabel("Várias imagens encontradas. Selecione a que deseja usar como logo:"))
                
                button_group = QtWidgets.QButtonGroup(dialog)
                
                # Grid para mostrar as imagens
                grid = QtWidgets.QGridLayout()
                rows = (len(images) // 3) + (1 if len(images) % 3 > 0 else 0)
                
                for i, img_path in enumerate(images):
                    try:
                        pixmap = QPixmap(img_path)
                        if pixmap.isNull():
                            continue
                            
                        # Container widget
                        container = QtWidgets.QWidget()
                        container_layout = QVBoxLayout(container)
                        
                        # Imagem
                        image_label = QLabel()
                        scaled_pixmap = pixmap.scaled(100, 100, Qt.AspectRatioMode.KeepAspectRatio)
                        image_label.setPixmap(scaled_pixmap)
                        image_label.setAlignment(Qt.AlignmentFlag.AlignCenter)
                        
                        # Radio button
                        radio = QtWidgets.QRadioButton(f"Imagem {i+1}")
                        radio.setProperty("img_path", img_path)
                        button_group.addButton(radio, i)
                        
                        container_layout.addWidget(image_label)
                        container_layout.addWidget(radio)
                        
                        row, col = i // 3, i % 3
                        grid.addWidget(container, row, col)
                        
                    except Exception as e:
                        print(f"Erro ao mostrar imagem {img_path}: {e}")
                
                layout.addLayout(grid)
                
                # Botões
                buttons = QtWidgets.QDialogButtonBox(
                    QtWidgets.QDialogButtonBox.StandardButton.Ok | 
                    QtWidgets.QDialogButtonBox.StandardButton.Cancel
                )
                buttons.accepted.connect(dialog.accept)
                buttons.rejected.connect(dialog.reject)
                layout.addWidget(buttons)
                
                dialog.setLayout(layout)
                
                if dialog.exec() == QtWidgets.QDialog.DialogCode.Accepted:
                    selected_button = button_group.checkedButton()
                    if selected_button:
                        selected_image = selected_button.property("img_path")
                    else:
                        self.progressBar.setValue(0)
                        return
                else:
                    self.progressBar.setValue(0)
                    return
            
            # Processar a imagem selecionada
            if selected_image and os.path.exists(selected_image):
                self.header_image_path = selected_image
                pixmap = QPixmap(selected_image)
                # Redimensionar para visualização
                scaled_pixmap = pixmap.scaled(80, 80, Qt.AspectRatioMode.KeepAspectRatio)
                self.image_preview.setPixmap(scaled_pixmap)
                self.selectedImageLabel.setText(f"Logo extraído de: {os.path.basename(file_path)}")
                self.progressBar.setValue(100)
                self.statusbar.showMessage("Logo importado com sucesso!", 3000)
                
                # Reset progress bar após 3 segundos
                QTimer.singleShot(3000, lambda: self.progressBar.setValue(0))
                
        except Exception as e:
            self.show_error("Erro", f"Erro ao importar logo: {str(e)}")
            self.progressBar.setValue(0)

    def on_import_header_complete(self):
        """Importar cabeçalho completo de um documento existente"""
        file_path, _ = QtWidgets.QFileDialog.getOpenFileName(
            self,
            "Selecionar Documento com Cabeçalho",
            "",
            "Documentos (*.pdf *.docx)"
        )
        
        if not file_path:
            return
            
        try:
            self.statusbar.showMessage(f"Importando cabeçalho de {os.path.basename(file_path)}...", 2000)
            self.progressBar.setValue(30)
            
            # Extrair informações do cabeçalho do documento
            header_info = {}
            logo_path = None
            
            # Processar com base no tipo de arquivo
            if file_path.endswith(".pdf"):
                # Extrair cabeçalho do PDF usando PyMuPDF
                import fitz  # PyMuPDF
                try:
                    doc = fitz.open(file_path)
                    page = doc[0]  # Primeira página geralmente contém o cabeçalho
                    
                    # Extrair logo/imagens
                    temp_dir = os.path.join(BASE_DIR, "output", "temp")
                    os.makedirs(temp_dir, exist_ok=True)
                    
                    # Extrair todas as imagens do começo da página
                    for img_index, img in enumerate(page.get_images(full=True)):
                        xref = img[0]
                        if xref:
                            pix = fitz.Pixmap(doc, xref)
                            if pix.width > 0 and pix.height > 0:
                                img_path = os.path.join(temp_dir, f"header_import_{img_index}.png")
                                if pix.alpha:
                                    pix = fitz.Pixmap(fitz.csRGB, pix)
                                pix.save(img_path)
                                if not logo_path:  # Usar primeira imagem como logo
                                    logo_path = img_path
                    
                    # Extrair texto da parte superior da página (primeiros 30% da altura)
                    rect = fitz.Rect(0, 0, page.rect.width, page.rect.height * 0.3)
                    header_text = page.get_text("text", clip=rect).strip()
                    
                    # Dividir em linhas e processar
                    lines = header_text.split('\n')
                    if lines:
                        # A primeira linha é geralmente o nome da escola
                        school_name = lines[0].strip()
                        
                        # Tentar identificar informações específicas nas linhas
                        subject_line = None
                        block_line = None
                        evaluation_type = None
                        teacher_name = None
                        
                        for line in lines[1:]:
                            line = line.strip()
                            if not line:
                                continue
                                
                            # Identificar linha de disciplina/professor
                            if "–" in line or "-" in line:
                                if any(subj in line.upper() for subj in [
                                    "HISTÓRIA", "MATEMÁTICA", "PORTUGUÊS", "GEOGRAFIA", 
                                    "CIÊNCIAS", "BIOLOGIA", "FÍSICA", "QUÍMICA", "INGLÊS"
                                ]):
                                    # Extrair disciplina e professor
                                    parts = re.split(r'[-–]\s*PROF\.?[º|O|ª|A]?\.?\s*', line, flags=re.IGNORECASE)
                                    if len(parts) >= 2:
                                        header_info['subject'] = parts[0].strip()
                                        header_info['teacher_info'] = parts[1].strip()
                                    else:
                                        header_info['subject'] = line
                            
                            # Identificar bloco/turma
                            elif "BLOCO" in line.upper() or "TURMA" in line.upper():
                                header_info['block_info'] = line
                            
                            # Identificar tipo de avaliação
                            elif "AVALIATIVA" in line.upper() or "PROVA" in line.upper() or "EXAME" in line.upper():
                                header_info['evaluation_type'] = line
                        
                        # Definir o nome da escola
                        header_info['school_name'] = school_name
                    
                    doc.close()
                    
                except Exception as e:
                    self.show_error("Erro", f"Erro ao extrair cabeçalho do PDF: {str(e)}")
                    self.progressBar.setValue(0)
                    return
                
            elif file_path.endswith(".docx"):
                # Extrair cabeçalho do DOCX
                import zipfile
                from docx import Document
                import re
                
                try:
                    doc = Document(file_path)
                    
                    # Extrair imagens do cabeçalho
                    temp_dir = os.path.join(BASE_DIR, "output", "temp")
                    os.makedirs(temp_dir, exist_ok=True)
                    
                    # Verificar se há imagens no cabeçalho
                    for section in doc.sections:
                        if section.header:
                            header_part = section.header.part
                            for rel in header_part.rels.values():
                                if "image" in rel.reltype:
                                    image_name = rel.target_ref.split('/')[-1]
                                    
                                    # Abrir o arquivo DOCX como ZIP e extrair a imagem
                                    zip_file = zipfile.ZipFile(file_path)
                                    image_data = zip_file.read(f"word/{rel.target_ref}")
                                    img_path = os.path.join(temp_dir, f"docx_header_{image_name}")
                                    
                                    with open(img_path, "wb") as f:
                                        f.write(image_data)
                                    
                                    if not logo_path:  # Usar primeira imagem como logo
                                        logo_path = img_path
                    
                    # Extrair texto do cabeçalho
                    header_text = ""
                    for section in doc.sections:
                        if section.header:
                            for paragraph in section.header.paragraphs:
                                if paragraph.text.strip():
                                    header_text += paragraph.text.strip() + "\n"
                    
                    # Se não encontrou nada no cabeçalho, pegar os primeiros parágrafos do documento
                    if not header_text and doc.paragraphs:
                        # Considerar os primeiros parágrafos (até 6) como possível cabeçalho
                        for i, paragraph in doc.paragraphs[:6]:
                            if paragraph.text.strip():
                                header_text += paragraph.text.strip() + "\n"
                    
                    # Processar o texto do cabeçalho
                    if header_text:
                        lines = header_text.split('\n')
                        if lines:
                            # A primeira linha é geralmente o nome da escola
                            school_name = lines[0].strip()
                            
                            # Processar as linhas restantes
                            for line in lines[1:]:
                                line = line.strip()
                                if not line:
                                    continue
                                    
                                # Identificar linha de disciplina/professor
                                if "–" in line or "-" in line:
                                    if any(subj in line.upper() for subj in [
                                        "HISTÓRIA", "MATEMÁTICA", "PORTUGUÊS", "GEOGRAFIA", 
                                        "CIÊNCIAS", "BIOLOGIA", "FÍSICA", "QUÍMICA", "INGLÊS"
                                    ]):
                                        # Extrair disciplina e professor
                                        parts = re.split(r'[-–]\s*PROF\.?[º|O|ª|A]?\.?\s*', line, flags=re.IGNORECASE)
                                        if len(parts) >= 2:
                                            header_info['subject'] = parts[0].strip()
                                            header_info['teacher_info'] = parts[1].strip()
                                        else:
                                            header_info['subject'] = line
                                
                                # Identificar bloco/turma
                                elif "BLOCO" in line.upper() or "TURMA" in line.upper():
                                    header_info['block_info'] = line
                                
                                # Identificar tipo de avaliação
                                elif "AVALIATIVA" in line.upper() or "PROVA" in line.upper() or "EXAME" in line.upper():
                                    header_info['evaluation_type'] = line
                        
                        # Definir o nome da escola
                        header_info['school_name'] = school_name
                    
                except Exception as e:
                    self.show_error("Erro", f"Erro ao extrair cabeçalho do DOCX: {str(e)}")
                    self.progressBar.setValue(0)
                    return
            
            # Atualizar os campos do cabeçalho
            if header_info.get('school_name'):
                self.headerTextInput.setText(header_info['school_name'])
            if header_info.get('subject'):
                self.subjectInput.setText(header_info['subject'])
            if header_info.get('teacher_info'):
                self.teacherInput.setText(header_info['teacher_info'])
            if header_info.get('block_info'):
                self.blockInput.setText(header_info['block_info'])
            if header_info.get('evaluation_type'):
                self.evaluationInput.setText(header_info['evaluation_type'])
            
            # Atualizar a imagem do cabeçalho
            if logo_path and os.path.exists(logo_path):
                self.header_image_path = logo_path
                pixmap = QPixmap(logo_path)
                # Redimensionar para visualização
                scaled_pixmap = pixmap.scaled(80, 80, Qt.AspectRatioMode.KeepAspectRatio)
                self.image_preview.setPixmap(scaled_pixmap)
                self.selectedImageLabel.setText(f"Logo importado de: {os.path.basename(file_path)}")
            
            self.progressBar.setValue(100)
            self.statusbar.showMessage("Cabeçalho importado com sucesso!", 3000)
            
            # Reset progress bar após 3 segundos
            QTimer.singleShot(3000, lambda: self.progressBar.setValue(0))
            
        except Exception as e:
            self.show_error("Erro", f"Erro ao importar cabeçalho: {str(e)}")
            self.progressBar.setValue(0)



